from flask import Flask, render_template, request, send_file
from werkzeug.utils import secure_filename
import os
from docx import Document
from docx2pdf import convert as docx2pdf_convert
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
from pdf2docx import Converter

app = Flask(__name__)
UPLOAD_FOLDER = "uploads"
OUTPUT_FOLDER = "converted"
os.makedirs(UPLOAD_FOLDER, exist_ok=True)
os.makedirs(OUTPUT_FOLDER, exist_ok=True)
app.config['UPLOAD_FOLDER'] = UPLOAD_FOLDER
app.config['OUTPUT_FOLDER'] = OUTPUT_FOLDER

# Conversion functions
def txt_to_pdf(txt_file, output_pdf):
    c = canvas.Canvas(output_pdf, pagesize=letter)
    width, height = letter
    with open(txt_file, "r", encoding="utf-8") as f:
        lines = f.readlines()
    y = height - 40
    for line in lines:
        c.drawString(40, y, line.strip())
        y -= 15
        if y < 40:
            c.showPage()
            y = height - 40
    c.save()

def txt_to_docx(txt_file, output_docx):
    doc = Document()
    with open(txt_file, "r", encoding="utf-8") as f:
        for line in f:
            doc.add_paragraph(line.strip())
    doc.save(output_docx)

def pdf_to_docx(pdf_file, output_docx):
    cv = Converter(pdf_file)
    cv.convert(output_docx, start=0, end=None)
    cv.close()

def convert_file(file_path, output_type):
    name, ext = os.path.splitext(file_path)
    ext = ext.lower()
    base_name = os.path.basename(name)
    output_file = os.path.join(app.config['OUTPUT_FOLDER'], f"{base_name}.{output_type}")

    if ext == ".txt":
        if output_type == "pdf":
            txt_to_pdf(file_path, output_file)
        elif output_type in ["docx", "doc"]:
            txt_to_docx(file_path, output_file)
        elif output_type == "ps":
            txt_to_pdf(file_path, output_file)  # PS can be done via PDF
    elif ext == ".docx":
        if output_type == "pdf":
            docx2pdf_convert(file_path, output_file)
        else:
            doc = Document(file_path)
            doc.save(output_file)
    elif ext == ".pdf":
        if output_type in ["docx", "doc"]:
            pdf_to_docx(file_path, output_file)
        elif output_type == "pdf":
            # PDF to PDF, just copy
            import shutil
            shutil.copy(file_path, output_file)
    else:
        return None
    return output_file

@app.route("/", methods=["GET", "POST"])
def index():
    if request.method == "POST":
        uploaded_file = request.files["file"]
        output_type = request.form.get("output_type")
        if uploaded_file.filename != "":
            filename = secure_filename(uploaded_file.filename)
            file_path = os.path.join(app.config['UPLOAD_FOLDER'], filename)
            uploaded_file.save(file_path)
            output_file = convert_file(file_path, output_type)
            if output_file:
                return send_file(output_file, as_attachment=True)
            else:
                return f"Conversion to {output_type.upper()} not supported."
    return render_template("index.html")

if __name__ == "__main__":
    app.run(debug=True)
